"""
Update 论坛详细设计3.0.docx with new AI and advanced features.
Fixed version - correct tuple order (text, style_name).
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

SRC = r'C:\Users\Administrator\Desktop\论坛详细设计3.0.docx'
DST = r'C:\Users\Administrator\Desktop\论坛详细设计3.0_完善版.docx'

doc = Document(SRC)

def make_paragraph(text, heading_level=None):
    """Create a properly formatted paragraph element.
    heading_level: 1 for Heading1, 2 for Heading2, 3 for Heading3, 4 for Heading4
    """
    new_p = OxmlElement('w:p')
    new_pPr = OxmlElement('w:pPr')

    # This document uses numeric style IDs: 2=Heading1, 3=Heading2, 4=Heading3, 5=Heading4
    style_map = {
        1: '2', 2: '3', 3: '4', 4: '5'
    }
    font_size_map = {
        1: '32', 2: '28', 3: '26', 4: '24'
    }

    if heading_level and heading_level in style_map:
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), style_map[heading_level])
        new_pPr.append(pStyle)
        font_size = font_size_map[heading_level]
        is_bold = True
    else:
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), 'Normal')
        new_pPr.append(pStyle)
        font_size = '24'
        is_bold = False

    # Set spacing for headings to match original
    if heading_level == 2:
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '260')
        spacing.set(qn('w:after'), '260')
        spacing.set(qn('w:line'), '416')
        spacing.set(qn('w:lineRule'), 'auto')
        new_pPr.append(spacing)

    if heading_level == 3:
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '260')
        spacing.set(qn('w:after'), '260')
        spacing.set(qn('w:line'), '416')
        spacing.set(qn('w:lineRule'), 'auto')
        new_pPr.append(spacing)

    new_p.append(new_pPr)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rFonts = OxmlElement('w:rFonts')
    if heading_level:
        rFonts.set(qn('w:eastAsia'), '黑体')
    else:
        rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), font_size)
    rPr.append(sz)

    if is_bold:
        b = OxmlElement('w:b')
        rPr.append(b)

    r.append(rPr)

    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    new_p.append(r)

    return new_p

def insert_before(ref_idx, paragraphs_data):
    """Insert paragraphs before the given paragraph index.
    paragraphs_data is a list of (text, heading_level) tuples.
    heading_level: None for Normal, 1/2/3/4 for headings.
    """
    ref_para = doc.paragraphs[ref_idx]
    ref_elem = ref_para._element
    # addprevious() makes each new element the immediate previous sibling of ref_elem.
    # So iterating in forward order (A, B, C) gives A, B, C, ref — correct.
    for text, h_level in paragraphs_data:
        new_p = make_paragraph(text, h_level)
        ref_elem.addprevious(new_p)

# ============ Build all new content ============

# --- PART 1: Update 1.3 ---
idx_tech_stack = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '技术栈说明：':
        idx_tech_stack = i
        break

if idx_tech_stack:
    modules_13 = [
        ("AI多智能体学习伴侣模块：AI规划任务、答疑Agent（RAG+Tool Calling）、监督Agent干预、心理Agent情绪识别、复盘Agent周报、行为分析Agent。", None),
        ("AI模拟复试官模块：多模态面试会话（文字+语音+TTS）、按院校风格追问、面试评估报告（内容深度/语言表达/心理状态/改进建议）。", None),
        ("上岸认证与经验贴模块：上岸用户录取通知书认证、结构化经验贴发布、按本科+目标院校精准检索相似上岸路径。", None),
        ("学习小组群聊模块：创建/加入学习小组、组内WebSocket实时群聊、文件分享、组长管理成员与公告。", None),
        ("会员增值服务模块：免费/VIP分级套餐、AI功能每日配额管理、Redis预扣+MySQL持久化双重保障。", None),
        ("AI智能择校引擎模块：基于用户画像+上岸数据库的院校推荐（保底/合适/冲刺三档）、相似上岸者案例匹配。", None),
        ("OCR智能错题本模块：拍照OCR识别→知识点定位→入错题本→艾宾浩斯遗忘曲线自动推送复习→PDF导出。", None),
        ("内容推荐与热度算法模块：用户画像推荐知识点、帖子热度衰减算法（Hacker News）、新帖冷启动曝光。", None),
        ("安全防护增强模块：图形验证码/滑块验证、接口限流（Rate Limiting）、敏感词库、XSS清洗、防SQL注入、全站HTTPS。", None),
    ]
    insert_before(idx_tech_stack, modules_13)
    print("✅ Part 1: Updated 1.3 with new modules")

# Find updated indices after insertion
def find_heading(text):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == text:
            return i
    return None

idx_db = find_heading('七、数据库详细设计')
idx_73 = find_heading('7.3 数据库索引设计')
idx_9 = find_heading('九、前后端对接规范')
idx_13_test = find_heading('十三、测试设计')
idx_12_4 = find_heading('12.4 学习打卡与积分奖励流程')
idx_frontend = find_heading('前端性能优化与已知问题')

print(f"New indices: db={idx_db}, 7.3={idx_73}, 九={idx_9}, 十三={idx_13_test}, 12.4={idx_12_4}, frontend={idx_frontend}")

# --- PART 2: Module sections 6.8-6.15 ---
if idx_db:
    new_modules = [
        ("6.8 AI多智能体学习伴侣模块(AI Agent Module)", 2),
        ("6.8.1 模块概述", 3),
        ("AI多智能体学习伴侣是本项目的核心创新模块，由6个专职Agent组成智能体团队：规划Agent（PlannerAgent）根据用户目标院校/当前进度/剩余时间动态生成每日任务清单；答疑Agent（TutorAgent）基于RAG知识库+Tool Calling定位考点链路；监督Agent（SupervisorAgent）连续打卡断签或学习时长下滑时主动介入发送Push；心理Agent（PsychologyAgent）识别打卡日记的负面情绪给出安抚引导；复盘Agent（ReviewAgent）每周自动生成学情透视周报定位薄弱点；行为分析Agent（BehaviorAnalysisAgent）基于用户浏览/点赞/收藏/搜索行为构建认知画像。", None),
        ("6.8.2 核心Agent说明", 3),
        ("（1）规划Agent（PlannerAgent）：每日根据用户连续打卡天数、累计学习时长、认知画像生成3条定制任务（HIGH/MEDIUM/LOW优先级），附带智能体叮嘱。通过事件驱动（UserCheckInEvent）在用户打卡后自动触发规划，异步执行不阻塞打卡主流程。", None),
        ("（2）答疑Agent（TutorAgent）：双模式运行——Tool Calling模式（默认）和RAG模式（降级）。Tool Calling模式下LLM自主决定是否调用search_knowledge工具检索知识库，支持多轮工具调用（最多3轮）。支持图片多模态输入+OCR降级双轨制：优先尝试多模态Vision直接理解题目图片；失败则自动降级为PaddleOCR识别文字后文本回答。流式SSE输出token即到即推，前端无需等待完整响应。", None),
        ("（3）监督Agent（SupervisorAgent）：定时扫描（可配置cron表达式）全站用户学习状态，检测连续断签≥2天或近7天学习时长下降≥30%的用户，自动生成干预消息存入ai_intervention_log表。支持手动触发（路演演示用）。", None),
        ("（4）心理Agent（PsychologyAgent）：监听打卡日记创建事件（UserDiaryCreatedEvent），对日记内容进行NLP情感分析，识别焦虑/抑郁/自我否定等负面信号。分析结果异步写入ai_intervention_log，触发推送引导。", None),
        ("（5）复盘Agent（ReviewAgent）：每周自动生成Markdown格式学情透视周报，包含：本周学习总时长/日均时长、完成任务数/完成率、打卡连续天数、薄弱知识点TOP5（从错题本统计）、各科目学习时间分布、下周学习建议。周报持久化存储于ai_report表，支持历史查询。", None),
        ("（6）行为分析Agent（BehaviorAnalysisAgent）：监听用户行为事件（VIEW_POST/COLLECT_POST/SEARCH/LIKE_POST），增量更新用户认知画像（UserAiProfile.cognitiveProfile JSON字段），提取兴趣关键词和知识薄弱点，为推荐引擎提供数据基础。", None),
        ("6.8.3 核心数据表", 3),
        ("ai_daily_task：AI每日任务表（user_id, task_date, task_content, importance, tips, status[0未完成/1已完成]）", None),
        ("ai_intervention_log：AI干预日志表（user_id, agent_type, intervention_type, message, user_reaction[UNREAD/READ/DISMISS]）", None),
        ("ai_report：AI周报表（user_id, week_start, week_end, markdown内容）", None),
        ("ai_chat_session：AI对话会话表（user_id, title, is_deleted, created_at, updated_at）", None),
        ("ai_chat_message：AI对话消息表（session_id, role[user/assistant], content, image_url）", None),
        ("ai_knowledge_point：考研知识库表（subject, chapter, title, content, keywords, importance[HIGH/MEDIUM/LOW]）", None),
        ("ai_user_event：用户行为事件表（user_id, event_type[VIEW_POST/COLLECT_POST/SEARCH/LIKE_POST], event_data JSON）", None),
        ("user_ai_profile：用户AI画像表（user_id, cognitive_profile JSON, created_at, updated_at）", None),
        ("6.8.4 核心接口", 3),
        ("GET /api/ai/tasks — 获取今日AI规划任务", None),
        ("POST /api/ai/tasks/{taskId}/complete — 完成某条AI任务（触发TaskCompletedEvent）", None),
        ("GET /api/ai/interventions — 获取未读AI干预日志", None),
        ("PUT /api/ai/interventions/{id}/read — 标记干预日志已读", None),
        ("GET /api/ai/report — 获取本周AI学情透视周报", None),
        ("POST /api/ai/ask — 向答疑Agent提问（RAG+Tool Calling增强，支持图片URL）", None),
        ("POST /api/ai/ask/stream — 向答疑Agent流式提问（SSE，token级实时推送）", None),
        ("GET /api/ai/knowledge — 搜索知识点", None),
        ("GET /api/ai/chat/sessions — 获取AI对话会话列表", None),
        ("GET /api/ai/chat/sessions/{id}/messages — 获取某会话的消息列表", None),
        ("DELETE /api/ai/chat/sessions/{id} — 删除会话", None),
        ("GET /api/ai/summary — 获取社区首页AI摘要数据", None),
        ("GET /api/ai/report/history — 获取历史周报列表", None),
        ("POST /api/ai/events — 上报用户行为事件（浏览/收藏/搜索/点赞）", None),
        ("GET /api/ai/recommendations — 基于用户画像推荐知识点", None),
        ("POST /api/ai/agent/supervisor/trigger — 手动触发监督Agent（路演演示）", None),

        ("6.9 AI模拟复试官模块(Interview Module)", 2),
        ("6.9.1 模块概述", 3),
        ("AI模拟复试官是本项目的核心创新模块（C项），聚焦考研复试这一最焦虑、信息最不对称的环节。支持文字模式和语音模式，可指定目标院校风格进行差异化追问，最终输出多维度评估报告（内容深度/语言表达/心理状态/改进建议）。", None),
        ("6.9.2 核心功能", 3),
        ("（1）会话管理：创建面试会话时指定目标院校(targetSchool)、目标专业(targetMajor)、面试类型(interviewType: 中文面/英文面/综合面/压力面)，系统根据院校风格差异生成不同的提问策略（如清华偏重科研潜力、普通985偏重基础扎实度）。", None),
        ("（2）AI追问引擎：采用LLM角色扮演模式，模拟真实复试官的语气和追问逻辑。每轮对话包含：AI提问→用户回答→AI根据回答内容生成针对性追问，形成连贯的面试对话链。支持语音回答时长(speechDuration)和视频仪态(demeanor)数据采集。", None),
        ("（3）TTS语音合成：AI面试官的文本回复可通过TTS接口合成为MP3语音，支持中文/英文自动切换（根据面试类型），模拟真人面试场景。", None),
        ("（4）评估报告：结束面试后自动生成多维度评估报告（ReportVO），包含：内容深度评分（专业知识掌握程度）、语言表达评分（流畅度/逻辑性）、心理状态评估（紧张度/自信度）、改进建议列表。视频模式下额外包含仪态分析。", None),
        ("6.9.3 核心数据表", 3),
        ("interview_session：面试会话表（user_id, target_school, target_major, interview_type, status[IN_PROGRESS/FINISHED]）", None),
        ("interview_record：面试对话记录表（session_id, role[interviewer/candidate], content, speech_duration, demeanor_analysis JSON）", None),
        ("interview_report：面试评估报告表（session_id, content_score, expression_score, psychology_score, suggestions JSON, overall_report）", None),
        ("6.9.4 核心接口", 3),
        ("POST /api/interview/session/create — 创建面试会话（需指定院校/专业/面试类型）", None),
        ("GET /api/interview/session/{sessionId} — 获取会话信息", None),
        ("GET /api/interview/session/{sessionId}/records — 获取历史对话记录", None),
        ("POST /api/interview/session/{sessionId}/next-question — 发送回答并获取AI追问", None),
        ("POST /api/interview/session/{sessionId}/finish — 结束面试并生成评估报告", None),
        ("POST /api/interview/tts — TTS语音合成（文本→MP3音频）", None),

        ("6.10 上岸认证与经验贴模块(Certification & Experience Module)", 2),
        ("6.10.1 模块概述", 3),
        ("上岸认证是本项目的核心创新模块（D项），通过录取通知书/学信网截图OCR+人工审核双重认证机制建立可信用户体系。认证用户的经验贴打\"✅已认证\"金标，结构化字段支持精准检索（本科背景/跨考否/二战否/初试分/复试分/各科分数/用书/时间线），帮助学弟学妹按\"我的本科+目标院校+目标分数\"精准检索相似上岸路径。", None),
        ("6.10.2 核心功能", 3),
        ("（1）上岸认证：用户提交录取通知书截图+学信网截图→上传至OSS→系统OCR辅助识别关键信息→管理员审核（APPROVED/REJECTED）→认证通过后用户获得\"已认证\"标识。认证信息结构化存储：本科院校/专业、是否跨考、是否二战、初试总分/各科分数、复试分数、备考时间线、使用书籍。", None),
        ("（2）经验贴管理：认证用户可发布结构化经验贴，含标题、正文（Markdown）、本科背景、跨考标志、初试分/复试分、各科分数、备考用书、时间线等字段。支持点赞/收藏互动。非认证用户可浏览但不可发布经验贴。", None),
        ("（3）精准检索：用户可按目标院校+目标专业+本科院校层次+跨考/非跨考等条件筛选经验贴，找到\"与自己最相似的上岸路径\"。为AI择校引擎提供数据支撑。", None),
        ("6.10.3 核心数据表", 3),
        ("user_verification：用户认证表（user_id, real_name, id_card_masked, admission_school, admission_major, admission_proof_url, xuexin_proof_url, status[PENDING/APPROVED/REJECTED], reviewer_id, review_comment）", None),
        ("experience_post：经验贴表（user_id, title, content_md, undergrad_school, undergrad_major, is_cross_major, is_second_war, initial_total_score, reexam_score, subject_scores JSON, books_used, timeline JSON, is_verified, like_count, collect_count, view_count）", None),
        ("experience_like：经验贴点赞表（user_id, post_id）", None),
        ("experience_collect：经验贴收藏表（user_id, post_id）", None),
        ("6.10.4 核心接口", 3),
        ("POST /api/verification/submit — 提交认证申请", None),
        ("GET /api/verification/status — 查询认证状态", None),
        ("POST /api/experience/posts — 发布经验贴", None),
        ("GET /api/experience/posts — 分页查询经验贴（支持多维度筛选）", None),
        ("GET /api/experience/posts/{id} — 获取经验贴详情", None),
        ("POST /api/experience/like/{postId} — 点赞/取消点赞经验贴", None),
        ("POST /api/experience/collect/{postId} — 收藏/取消收藏经验贴", None),

        ("6.11 学习小组群聊模块(Chat Group Module)", 2),
        ("6.11.1 模块概述", 3),
        ("学习小组群聊模块允许用户创建/加入学习小组（如\"408刷题组\"、\"英语每日打卡组\"），组内支持WebSocket实时群聊消息、文件分享，组长可管理成员、设置小组公告。技术选型采用WebSocket实时通信，消息持久化到MySQL。", None),
        ("6.11.2 核心功能", 3),
        ("（1）小组管理：创建小组（名称、描述、最大成员数、是否公开）、编辑小组信息、解散小组（仅组长）。公开小组可被搜索和申请加入，私密小组需邀请码。", None),
        ("（2）成员管理：申请加入→组长/管理员审批→入组。组长可设置管理员、移除成员、禁言成员。成员可主动退组。", None),
        ("（3）实时群聊：基于WebSocket的实时消息推送，消息类型支持文本和文件分享。消息持久化存储，新成员入组可查看历史消息。未读消息计数。", None),
        ("（4）小组公告：组长/管理员可编辑小组置顶公告，全体成员可见。", None),
        ("6.11.3 核心数据表", 3),
        ("chat_group：小组表（name, description, owner_id, max_members, is_public, invite_code, announcement, status）", None),
        ("group_member：小组成员表（group_id, user_id, role[OWNER/ADMIN/MEMBER], status[PENDING/APPROVED/BANNED], joined_at）", None),
        ("group_message：群聊消息表（group_id, sender_id, message_type[TEXT/FILE], content, file_url, file_name）", None),
        ("6.11.4 核心接口", 3),
        ("POST /api/chat/groups — 创建学习小组", None),
        ("GET /api/chat/groups — 分页搜索小组列表", None),
        ("GET /api/chat/groups/{id} — 获取小组详情", None),
        ("POST /api/chat/groups/{id}/join — 申请加入小组", None),
        ("PUT /api/chat/groups/{id}/members/{userId}/approve — 审批入组申请", None),
        ("GET /api/chat/groups/{id}/messages — 获取群聊历史消息", None),
        ("WebSocket /ws/chat/group/{groupId} — 群聊实时通信", None),

        ("6.12 会员增值服务模块(Membership Module)", 2),
        ("6.12.1 模块概述", 3),
        ("会员增值服务模块提供免费/VIP分级套餐体系，通过AI功能每日配额管理实现增值变现。基于Redis Lua脚本原子预扣+MySQL持久化双写保障，确保配额准确性。采用@MembershipRequired注解声明式权限控制，零侵入接入业务代码。", None),
        ("6.12.2 核心功能", 3),
        ("（1）套餐管理：预定义3种套餐——免费版（基础功能，AI每日限3次）、VIP月卡（¥29.9/月，AI每日100次+模拟复试+择校推荐+TTS语音）、VIP年卡（¥199/年，全部功能无限次+专属客服）。套餐信息存储在membership_plan表，支持动态上下架。", None),
        ("（2）配额管理：AI功能（ai_ask/ocr/interview/school_recommend/interview_tts/weekly_report/ai_knowledge）采用Redis计数器+Lua脚本原子预扣+MySQL使用日志持久化的三层架构。每日0点通过定时任务自动重置配额。配额耗尽时返回402状态码，前端展示升级引导。", None),
        ("（3）订单系统：支持创建订单→支付回调→激活会员的完整流程。订单状态：PENDING→PAID→ACTIVATED。支持取消自动续费。", None),
        ("（4）声明式权限：@MembershipRequired(\"feature_key\")注解方法，通过AOP切面(MembershipAspect)自动校验会员权限和配额，未登录/配额不足/需VIP时抛出MembershipException，由全局异常处理器统一返回。", None),
        ("6.12.3 核心数据表", 3),
        ("membership_plan：套餐定义表（plan_code, plan_name, price, duration_days, features JSON, daily_quotas JSON, is_active）", None),
        ("user_membership：用户会员表（user_id, plan_id, status[ACTIVE/EXPIRED/CANCELLED], start_date, end_date, auto_renew）", None),
        ("membership_order：会员订单表（user_id, plan_id, order_no, amount, payment_status[PENDING/PAID/CANCELLED], payment_time）", None),
        ("user_usage_log：使用日志表（user_id, feature_key, used_at, membership_id）——MySQL持久化层", None),
        ("6.12.4 核心接口", 3),
        ("GET /api/membership/plans — 获取所有可选套餐", None),
        ("GET /api/membership/me — 获取当前用户会员状态与配额", None),
        ("GET /api/membership/check/{featureKey} — 检查某功能是否可用", None),
        ("POST /api/membership/upgrade — 升级套餐（创建订单）", None),
        ("POST /api/membership/cancel — 取消自动续费", None),

        ("6.13 AI智能择校引擎模块(School Select Module)", 2),
        ("6.13.1 模块概述", 3),
        ("AI智能择校引擎是本项目的核心创新模块（G项），区别于传统\"院校信息库\"（告诉你\"有哪些学校\"），择校引擎基于用户画像+真实上岸数据回答\"我能去哪\"。输入：本科院校/GPA/英语等级/备考时长/模考分数/风险偏好→输出：保底/合适/冲刺三档院校推荐+\"与你相似的N位上岸者去了哪些学校\"。形成\"用户越多→数据越准→吸引新用户\"的数据飞轮。", None),
        ("6.13.2 核心功能", 3),
        ("（1）院校推荐：用户填写个人画像（本科院校层次/GPA/英语等级/目标专业/备考时长/模考分数/风险偏好），匹配引擎(MatchEngineService)计算与历史上岸记录的相似度，输出保底/合适/冲刺三档院校推荐列表，每所院校附带匹配度评分和推荐理由。", None),
        ("（2）相似案例：展示与用户画像最相似的N位已上岸用户案例（脱敏后），包含：本科背景、上岸院校专业、初试分数、备考时间线。提供\"他们去了哪\"的可视化分布。", None),
        ("（3）院校库：维护全国研究生招生院校信息（院校名称、层次[985/211/双一流/普通]、所在地、研究生院官网、历年分数线趋势）。支持关键词搜索。", None),
        ("（4）推荐历史：保存用户每次推荐请求和结果，支持回溯查看历史推荐。", None),
        ("6.13.3 核心数据表", 3),
        ("school_info：院校信息表（name, tier[985/211/DOUBLE_FIRST/REGULAR], province, city, website, score_line_trend JSON）", None),
        ("school_major：院校专业表（school_id, major_name, major_code, degree_type[MASTER/PHD], exam_subjects JSON, enrollment_count）", None),
        ("admission_record：上岸记录表（user_id, school_id, major_id, initial_score, reexam_score, total_score, undergrad_school, undergrad_major, is_cross_major, prep_duration_months, year）——数据来源于认证用户", None),
        ("recommendation_history：推荐历史表（user_id, request_params JSON, result JSON, created_at）", None),
        ("6.13.4 核心接口", 3),
        ("POST /api/school-select/recommend — 获取择校推荐（需用户画像参数）", None),
        ("GET /api/school-select/history — 获取历史推荐记录", None),
        ("GET /api/school-select/schools — 查询院校列表（支持关键词搜索）", None),

        ("6.14 OCR智能错题本模块(Mistake Note Module)", 2),
        ("6.14.1 模块概述", 3),
        ("OCR智能错题本是本项目的核心创新模块（J项），实现\"拍照→OCR识别→知识点定位→入错题本→艾宾浩斯遗忘曲线自动推送复习→PDF导出\"的完整闭环。与AI答疑Agent联动：在AI对话中可直接快速保存为错题（quick-save接口），支持去重检查。", None),
        ("6.14.2 核心功能", 3),
        ("（1）OCR识别：用户上传题目图片→调用PaddleOCR引擎识别文字→返回识别结果和建议的科目/知识点。支持中文数学公式的较好识别。可指定科目提示提升识别精度。", None),
        ("（2）错题管理：创建错题（含题干、答案、解析、科目、知识点标签、掌握程度[WEAK/FAIR/STRONG]）、编辑错题、删除错题、分页查询错题列表（支持科目/掌握程度/知识点筛选）。", None),
        ("（3）艾宾浩斯复习计划：基于艾宾浩斯遗忘曲线（1天/2天/4天/7天/15天/30天/60天/120天共8个阶段），创建错题时自动设置首次复习日期=明天（第0阶段）。每日自动推送当天该复习的错题列表。完成复习后标记该阶段已复习，自动计算下一阶段复习日期。支持复习统计：各阶段复习进度、掌握程度分布。", None),
        ("（4）PDF导出：支持将错题本导出为PDF文件（便于打印），可指定导出范围（全部/按科目/按日期范围）。", None),
        ("（5）AI对话快速收藏：在AI答疑对话中，用户可一键将AI的回复保存到错题本（quick-save接口），自动提取AI消息中的知识点标签，支持去重检查（通过chatMessageId去重）。", None),
        ("（6）Markdown渲染：支持将错题解析内容渲染为Markdown格式（含LaTeX数学公式），便于富文本展示。", None),
        ("6.14.3 核心数据表", 3),
        ("mistake_note：错题表（user_id, subject, question, answer, analysis, knowledge_points, mastery[WEAK/FAIR/STRONG], source_type[MANUAL/OCR/AI_CHAT], image_url, chat_message_id, ebbinghaus_stage, next_review_date, is_deleted）", None),
        ("review_log：复习日志表（note_id, stage, review_date, reviewed_at, is_completed）", None),
        ("daily_plan：每日复习计划表（user_id, plan_date, total_tasks, completed_tasks）", None),
        ("mistake_notification：错题通知表（user_id, notification_type[REVIEW_REMINDER/STAGE_ADVANCE], message, is_read, created_at）", None),
        ("6.14.4 核心接口", 3),
        ("POST /api/mistake/ocr — OCR识别题目图片（需先调用上传接口获取图片URL）", None),
        ("POST /api/mistake/notes — 创建错题（从OCR结果或手动输入）", None),
        ("POST /api/mistake/quick-save — 从AI对话快速收藏（支持去重）", None),
        ("POST /api/mistake/check-saved — 批量检查AI消息是否已收藏", None),
        ("GET /api/mistake/notes — 分页查询错题列表", None),
        ("GET /api/mistake/notes/{id} — 获取错题详情", None),
        ("PUT /api/mistake/notes/{id} — 编辑错题", None),
        ("DELETE /api/mistake/notes/{id} — 删除错题（逻辑删除）", None),
        ("GET /api/mistake/reviews/today — 获取今日待复习错题列表", None),
        ("POST /api/mistake/reviews/{noteId}/complete — 完成某错题的当前阶段复习", None),
        ("GET /api/mistake/reviews/stats — 获取艾宾浩斯复习统计", None),
        ("POST /api/mistake/pdf/export — 导出错题本PDF", None),
        ("POST /api/mistake/markdown/render — Markdown渲染", None),

        ("6.15 内容推荐与热度算法模块(Recommendation Module)", 2),
        ("6.15.1 模块概述", 3),
        ("内容推荐算法模块实现基于用户画像的知识点推荐和基于热度衰减算法的帖子排序。前者根据用户认知画像中的兴趣关键词推荐相关知识库知识点，后者采用类似Hacker News的热度衰减算法避免旧帖长期霸榜，同时给予新帖冷启动曝光权重。", None),
        ("6.15.2 核心功能", 3),
        ("（1）知识点推荐：基于用户行为（浏览/点赞/收藏/搜索）构建兴趣关键词画像→匹配知识库中的相关知识点→输出推荐列表并附推荐理由。推荐不足3条时补充高频考点作为兜底。", None),
        ("（2）帖子热度衰减：定时任务(HotPostScheduler)周期性更新帖子热度分数，算法综合考虑：点赞数(weight=3)、评论数(weight=5)、收藏数(weight=4)、发布时间衰减因子（越新权重越高）、新帖冷启动加权。首页帖子列表按热度降序排列，保证用户看到高质量+新鲜的混合内容。", None),
        ("6.15.3 核心接口", 3),
        ("GET /api/ai/recommendations — 基于用户画像推荐知识点", None),
        ("GET /api/posts/page?sort=hot — 按热度排序获取帖子列表", None),
    ]
    insert_before(idx_db, new_modules)
    print(f"✅ Part 2: Added modules 6.8-6.15 ({len(new_modules)} paragraphs)")

# --- PART 3: Database tables ---
idx_73 = find_heading('7.3 数据库索引设计')
if idx_73:
    new_tables = [
        ("7.2.9 AI多智能体相关表", 3),
        ("（1）ai_daily_task（AI每日任务表）：id(BIGINT PK), user_id(BIGINT FK→sys_user.id), task_date(DATE), task_content(VARCHAR 500), importance(VARCHAR 10), tips(VARCHAR 300), status(TINYINT 0未完成/1已完成), created_at(DATETIME)。索引：idx_user_date(user_id, task_date)。", None),
        ("（2）ai_intervention_log（AI干预日志表）：id(BIGINT PK), user_id(BIGINT FK), agent_type(VARCHAR 50), intervention_type(VARCHAR 50), message(TEXT), user_reaction(VARCHAR 20 UNREAD/READ/DISMISS), created_at(DATETIME)。索引：idx_user_reaction(user_id, user_reaction)。", None),
        ("（3）ai_chat_session（AI对话会话表）：id(BIGINT PK), user_id(BIGINT FK), title(VARCHAR 100), is_deleted(TINYINT 默认0), created_at(DATETIME), updated_at(DATETIME)。索引：idx_user_deleted(user_id, is_deleted)。", None),
        ("（4）ai_chat_message（AI对话消息表）：id(BIGINT PK), session_id(BIGINT FK→ai_chat_session.id), role(VARCHAR 20 user/assistant), content(MEDIUMTEXT), image_url(VARCHAR 500), created_at(DATETIME)。索引：idx_session(session_id)。", None),
        ("（5）ai_knowledge_point（考研知识库表）：id(BIGINT PK), subject(VARCHAR 50), chapter(VARCHAR 200), title(VARCHAR 300), content(TEXT), keywords(VARCHAR 500), importance(VARCHAR 10 HIGH/MEDIUM/LOW), created_at(DATETIME)。索引：idx_subject_importance(subject, importance)；FULLTEXT全文索引(title, content, keywords)。", None),

        ("7.2.10 面试模块相关表", 3),
        ("（1）interview_session（面试会话表）：id(BIGINT PK), user_id(BIGINT FK), target_school(VARCHAR 100), target_major(VARCHAR 100), interview_type(VARCHAR 50 中文面/英文面/综合面/压力面), status(VARCHAR 20 IN_PROGRESS/FINISHED), created_at(DATETIME)。", None),
        ("（2）interview_record（面试记录表）：id(BIGINT PK), session_id(BIGINT FK), role(VARCHAR 20 interviewer/candidate), content(TEXT), speech_duration(DECIMAL 10,2), demeanor_analysis(JSON), created_at(DATETIME)。", None),
        ("（3）interview_report（面试报告表）：id(BIGINT PK), session_id(BIGINT FK), content_score(DECIMAL 5,2), expression_score(DECIMAL 5,2), psychology_score(DECIMAL 5,2), suggestions(JSON), overall_report(TEXT), created_at(DATETIME)。", None),

        ("7.2.11 会员与认证相关表", 3),
        ("（1）membership_plan（套餐定义表）：id(BIGINT PK), plan_code(VARCHAR 50 free/vip_monthly/vip_yearly), plan_name(VARCHAR 100), price(DECIMAL 10,2), duration_days(INT), features(JSON), daily_quotas(JSON), is_active(TINYINT)。", None),
        ("（2）user_membership（用户会员表）：id(BIGINT PK), user_id(BIGINT FK), plan_id(BIGINT FK), status(VARCHAR 20 ACTIVE/EXPIRED/CANCELLED), start_date(DATE), end_date(DATE), auto_renew(TINYINT)。", None),
        ("（3）user_verification（用户认证表）：id(BIGINT PK), user_id(BIGINT FK), admission_school(VARCHAR 100), admission_major(VARCHAR 100), admission_proof_url(VARCHAR 500), xuexin_proof_url(VARCHAR 500), status(VARCHAR 20 PENDING/APPROVED/REJECTED), reviewer_id(BIGINT), review_comment(VARCHAR 500)。", None),

        ("7.2.12 错题本与择校相关表", 3),
        ("（1）mistake_note（错题表）：id(BIGINT PK), user_id(BIGINT FK), subject(VARCHAR 50), question(TEXT), answer(TEXT), analysis(TEXT), knowledge_points(VARCHAR 500), mastery(VARCHAR 20), ebbinghaus_stage(INT 0-7), next_review_date(DATE), is_deleted(TINYINT)。索引：idx_user_next_review(user_id, next_review_date, is_deleted)。", None),
        ("（2）school_info（院校信息表）：id(BIGINT PK), name(VARCHAR 200), tier(VARCHAR 50), province(VARCHAR 50), city(VARCHAR 50), website(VARCHAR 500), score_line_trend(JSON)。", None),
        ("（3）admission_record（上岸记录表）：id(BIGINT PK), user_id(BIGINT FK), school_id(BIGINT FK), major_id(BIGINT FK), initial_score(DECIMAL 5,1), reexam_score(DECIMAL 5,1), undergrad_school(VARCHAR 200), is_cross_major(TINYINT), prep_duration_months(INT), year(INT)。", None),

        ("7.2.13 经验贴与小组相关表", 3),
        ("（1）experience_post（经验贴表）：id(BIGINT PK), user_id(BIGINT FK), title(VARCHAR 200), content_md(MEDIUMTEXT), undergrad_school(VARCHAR 200), initial_total_score(DECIMAL 5,1), reexam_score(DECIMAL 5,1), subject_scores(JSON), is_verified(TINYINT), like_count(INT), collect_count(INT), view_count(INT)。", None),
        ("（2）chat_group（学习小组表）：id(BIGINT PK), name(VARCHAR 100), description(VARCHAR 500), owner_id(BIGINT FK), max_members(INT), is_public(TINYINT), announcement(VARCHAR 1000)。", None),
        ("（3）group_message（群聊消息表）：id(BIGINT PK), group_id(BIGINT FK), sender_id(BIGINT FK), message_type(VARCHAR 20 TEXT/FILE), content(TEXT), file_url(VARCHAR 500), created_at(DATETIME)。", None),
    ]
    insert_before(idx_73, new_tables)
    print(f"✅ Part 3: Added database tables ({len(new_tables)} paragraphs)")

# --- PART 4: API sections ---
idx_9 = find_heading('九、前后端对接规范')
if idx_9:
    new_apis = [
        ("8.8 AI Agent模块接口", 2),

        ("8.8.1 获取今日AI规划任务", 3),
        ("请求方式：GET /api/ai/tasks", None),
        ("认证要求：需登录（Sa-Token）", None),
        ("请求参数：无", None),
        ("返回数据：List<AiTaskVO>，包含任务ID、任务内容、重要性、智能体叮嘱、完成状态。", None),
        ("业务逻辑：从ai_daily_task表查询当前用户今日的所有规划任务，按创建时间升序排列。", None),

        ("8.8.2 向答疑Agent提问", 3),
        ("请求方式：POST /api/ai/ask", None),
        ("认证要求：需登录 + 会员权限@MembershipRequired(\"ai_ask\")", None),
        ("请求参数：question(必填, 问题文本), subject(选填, 限定学科), imageUrl(选填, 题目图片URL), sessionId(选填, 续接会话ID)。", None),
        ("返回数据：{answer, question, sessionId, sessionTitle}。", None),
        ("业务逻辑：自动创建/续接对话会话→保存用户消息→TutorAgent双轨制回答（多模态优先+OCR降级）→保存AI回复→返回。每次调用消耗1次AI配额。", None),

        ("8.8.3 流式答疑（SSE）", 3),
        ("请求方式：POST /api/ai/ask/stream（produces: text/event-stream）", None),
        ("认证要求：需登录（手动校验，不使用注解）", None),
        ("请求参数：同上", None),
        ("返回方式：SSE事件流，每收到一个token立即推送JSON: {\"content\":\"token_text\"}，首条事件为会话元数据{\"type\":\"meta\",\"sessionId\":123,\"title\":\"...\"}。错误通过error事件推送。", None),
        ("技术要点：异步执行避免阻塞请求线程；Redis Lua原子预扣配额；AI调用失败时自动退款配额。", None),

        ("8.8.4 获取AI周报", 3),
        ("请求方式：GET /api/ai/report", None),
        ("认证要求：需登录", None),
        ("返回数据：{markdown: \"周报Markdown内容\"}，由复盘Agent实时生成。", None),

        ("8.9 面试模块接口", 2),

        ("8.9.1 创建面试会话", 3),
        ("请求方式：POST /api/interview/session/create", None),
        ("认证要求：需登录 + 会员权限@MembershipRequired(\"interview\")", None),
        ("请求参数：targetSchool(目标院校), targetMajor(目标专业), interviewType(面试类型: 中文面/英文面/综合面/压力面)。", None),
        ("返回数据：完整的InterviewSession对象，含sessionId和status=IN_PROGRESS。", None),

        ("8.9.2 发送回答获取AI追问", 3),
        ("请求方式：POST /api/interview/session/{sessionId}/next-question", None),
        ("请求参数：answer(用户回答文本), speechDuration(语音时长，语音模式), demeanor(仪态数据，视频模式)。", None),
        ("返回数据：InterviewRecord对象，包含AI追问内容和角色标记。", None),

        ("8.9.3 结束面试生成报告", 3),
        ("请求方式：POST /api/interview/session/{sessionId}/finish", None),
        ("返回数据：ReportVO对象，包含内容深度/语言表达/心理状态评分和改进建议。", None),

        ("8.9.4 TTS语音合成", 3),
        ("请求方式：POST /api/interview/tts", None),
        ("认证要求：需登录 + 会员权限@MembershipRequired(\"interview_tts\")", None),
        ("请求参数：text(待合成文本), sessionId(选填，用于自动判断中英文语音类型)。", None),
        ("返回数据：audio/mpeg二进制音频流。", None),

        ("8.10 会员模块接口", 2),
        ("8.10.1 获取套餐列表", 3),
        ("请求方式：GET /api/membership/plans（无需登录）", None),
        ("返回数据：List<MembershipPlanVO>，含套餐名称、价格、时长、功能列表、每日配额。", None),

        ("8.10.2 检查功能可用性", 3),
        ("请求方式：GET /api/membership/check/{featureKey}", None),
        ("返回数据：{featureKey, available, used, limit, remaining, reason}。reason枚举：OK/VIP_REQUIRED/QUOTA_EXHAUSTED。", None),

        ("8.11 错题本模块接口", 2),
        ("8.11.1 OCR识别", 3),
        ("请求方式：POST /api/mistake/ocr", None),
        ("认证要求：需登录 + 会员权限@MembershipRequired(\"ocr\")", None),
        ("请求参数：imageUrl(必填，图片URL), subject(选填，科目提示)。", None),
        ("返回数据：OCRResultVO，含识别的文本内容和建议的科目/知识点。", None),

        ("8.11.2 今日待复习", 3),
        ("请求方式：GET /api/mistake/reviews/today", None),
        ("返回数据：List<ReviewTaskVO>，按艾宾浩斯阶段分组展示今日该复习的错题。", None),

        ("8.11.3 导出PDF", 3),
        ("请求方式：POST /api/mistake/pdf/export", None),
        ("请求参数：subject(选填，限定科目), startDate/endDate(选填，日期范围)。", None),
        ("返回数据：PDF文件二进制流。", None),

        ("8.12 择校引擎接口", 2),
        ("8.12.1 获取择校推荐", 3),
        ("请求方式：POST /api/school-select/recommend", None),
        ("认证要求：需登录 + 会员权限@MembershipRequired(\"school_recommend\")", None),
        ("请求参数：undergradSchool(本科院校), gpa, englishLevel, targetMajor, prepDuration, mockScore, riskPreference(风险偏好: SAFE/BALANCED/AGGRESSIVE)。", None),
        ("返回数据：RecommendationResultVO，包含保底/合适/冲刺三档院校列表+相似上岸者案例。", None),

        ("8.13 经验贴模块接口", 2),
        ("8.13.1 分页查询经验贴", 3),
        ("请求方式：GET /api/experience/posts", None),
        ("请求参数：page, size, undergradSchool(选填), targetSchool(选填), isCrossMajor(选填), isVerified(选填), sortBy(排序: latest/hottest)。", None),
        ("返回数据：PageInfo<ExperiencePostVO>，认证用户的经验贴含\"✅已认证\"金标。", None),
    ]
    insert_before(idx_9, new_apis)
    print(f"✅ Part 4: Added API sections 8.8-8.13 ({len(new_apis)} paragraphs)")

# --- PART 5: Flow charts ---
idx_13_test = find_heading('十三、测试设计')
if idx_13_test:
    new_flows = [
        ("12.5 AI答疑RAG+Tool Calling流程", 2),
        ("用户提问→提取关键词→搜索知识库→构建上下文→LLM判断是否需要工具调用→如需则调用search_knowledge→获取知识库结果→LLM综合生成回答→保存对话记录→返回答案。若Tool Calling失败则降级为纯RAG模式。", None),

        ("12.6 AI模拟复试流程", 2),
        ("创建面试会话（指定院校/专业/类型）→AI生成开场白→用户回答→AI追问（根据回答内容+院校风格）→循环多轮→用户结束面试→生成多维度评估报告（内容/表达/心理/改进建议）→持久化存储面试记录。", None),

        ("12.7 OCR错题本闭环流程", 2),
        ("拍照上传→调用PaddleOCR识别→返回识别文字+知识点建议→用户确认/编辑→保存到错题本→自动计算艾宾浩斯第0阶段（明日复习）→每日自动推送待复习错题→完成复习→自动推进到下一阶段→8阶段完成后标记为已掌握。", None),

        ("12.8 会员配额消费流程", 2),
        ("用户请求AI功能→@MembershipRequired AOP拦截→校验登录状态→查询会员等级→Redis Lua脚本原子检查+预扣配额→配额足够则放行→执行业务逻辑→MySQL持久化使用记录→配额不足则返回402→前端展示升级引导。AI调用失败时自动退款配额。", None),

        ("12.9 上岸认证+经验贴流程", 2),
        ("用户提交录取通知书/学信网截图→上传至OSS→管理员审核（OCR辅助）→审核通过→用户获得\"已认证\"标识→发布结构化经验贴→系统打\"✅已认证\"金标→其他用户按条件精准检索→择校引擎引用数据→数据飞轮形成。", None),
    ]
    insert_before(idx_13_test, new_flows)
    print(f"✅ Part 5: Added flow sections 12.5-12.9 ({len(new_flows)} paragraphs)")

# --- PART 6: Extend 15.2 ---
idx_frontend = find_heading('前端性能优化与已知问题')
if idx_frontend:
    new_extend = [
        ("AI多智能体扩展：引入LangChain4j/Spring AI框架替代自研Agent调度，支持更复杂的多Agent协作和上下文管理。接入向量数据库（Milvus/pgvector）实现长期记忆和语义检索增强RAG效果。", None),
        ("知识图谱建设：引入Neo4j图数据库构建考研知识图谱，将考研大纲→知识点→真题→解析→错题构建为图结构，支持前置依赖反向诊断和个性化学习路径推荐。", None),
        ("实时音视频面试：集成WebRTC（SRS/mediasoup）实现实时视频面试，结合OpenCV人脸表情识别分析考生仪态，提升模拟复试的真实感和评估维度。", None),
        ("推荐算法升级：引入协同过滤（Collaborative Filtering）和深度学习推荐模型，基于用户行为和相似用户群实现更精准的内容推荐。", None),
        ("移动端覆盖：优先开发微信小程序（与微信OAuth登录联动，分享裂变效果好），后续考虑响应式Web或H5+PWA作为补充。", None),
    ]
    insert_before(idx_frontend, new_extend)
    print(f"✅ Part 6: Updated 15.2 with suggestions ({len(new_extend)} paragraphs)")

# ============ Save ============
doc.save(DST)
print(f"\n✅ Document saved to: {DST}")
print(f"Total paragraphs: {len(doc.paragraphs)}")
